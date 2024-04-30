import os
import sys
import joblib
from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QPushButton, QTextEdit, QFileDialog, QLabel
from PyQt5.QtGui import QColor, QPalette
from PyQt5.QtCore import Qt
import docx
import pdfplumber
import sqlite3
from datetime import datetime, timedelta
import torch 
from transformers import RobertaTokenizer, RobertaForSequenceClassification, RobertaConfig
import pytesseract
from PIL import Image

class TextExtractorApp(QWidget):
    def __init__(self):
        super().__init__()
        self.model = joblib.load('random_forest_model.joblib')
        self.vectorizer = joblib.load('tfidf_vectorizer.joblib')
        current_directory = os.path.dirname(os.path.abspath(__file__))
        model_path = os.path.join(current_directory, 'roberta_classification_model')
        self.tokenizer = RobertaTokenizer.from_pretrained(model_path)
        self.number_to_category = {
            0: 'беспилотники',
            1: 'приказ',
            2: 'конфликт',
            3: 'ИБ',
            4: 'ракета',
            5: 'радиация',
            6: 'преступления',
            7: 'терроризм',
            8: 'разведка',
            9: 'украина'
        }
        self.roberta_number_to_category = {
                0: 'обычные',
                1: 'военные'
            }
        num_classes = len(self.roberta_number_to_category)
        self.roberta_model = RobertaForSequenceClassification.from_pretrained(model_path, num_labels=num_classes)
        self.initUI()
        self.create_database()

    def create_database(self):
        self.conn = sqlite3.connect('actions.db')
        self.c = self.conn.cursor()
        self.c.execute('''CREATE TABLE IF NOT EXISTS actions
                        (id INTEGER PRIMARY KEY AUTOINCREMENT, date TEXT, category TEXT, roberta_category TEXT, file_path TEXT)''')
        self.conn.commit()

    def closeEvent(self, event):
        self.conn.close()
        event.accept()

    def initUI(self):
        layout = QVBoxLayout(self)

        # Установка темной темы
        self.setStyleSheet("QWidget { background-color: #2B2B2B; color: white; }")

        self.btn_load = QPushButton('Загрузить файл', self)
        self.btn_load.setStyleSheet("""
            QPushButton {
                background-color: #00FF00;
                color: black;
                border-radius: 10px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #00E000;
            }
            QPushButton:pressed {
                background-color: #00C000;
            }
        """)
        self.btn_load.clicked.connect(self.loadFile)
        layout.addWidget(self.btn_load)

        self.btn_save = QPushButton('Сохранить файл', self)
        self.btn_save.setStyleSheet("""
            QPushButton {
                background-color: #00FF00;
                color: black;
                border-radius: 10px;
                padding: 5px;
            }
            QPushButton:hover {
                background-color: #00E000;
            }
            QPushButton:pressed {
                background-color: #00C000;
            }
        """)
        self.btn_save.clicked.connect(self.saveFile)
        self.btn_save.setEnabled(False)
        layout.addWidget(self.btn_save)

        self.btn_day = QPushButton('За день', self)
        self.btn_day.clicked.connect(self.show_actions_for_day)
        layout.addWidget(self.btn_day)

        self.btn_week = QPushButton('За неделю', self)
        self.btn_week.clicked.connect(self.show_actions_for_week)
        layout.addWidget(self.btn_week)

        self.btn_month = QPushButton('За месяц', self)
        self.btn_month.clicked.connect(self.show_actions_for_month)
        layout.addWidget(self.btn_month)

        self.text_edit = QTextEdit(self)
        self.text_edit.setReadOnly(True)
        layout.addWidget(self.text_edit)

        self.label_category = QLabel('Категория: Не определена', self)
        layout.addWidget(self.label_category)

        self.label_roberta_category = QLabel('Категория (RoBERTa): Не определена', self)
        layout.addWidget(self.label_roberta_category)

        self.setLayout(layout)
        self.setWindowTitle('Text Extractor with Classification')
        self.setGeometry(300, 300, 600, 400)

    def loadFile(self):
        options = QFileDialog.Options()
        filename, _ = QFileDialog.getOpenFileName(
            self, "Выберите файл", "", "Word Documents (*.docx);;PDF Files (*.pdf);;Image Files (*.png *.jpg *.jpeg)", options=options)
        if filename:
            self.displayContent(filename)
            self.btn_save.setEnabled(True)

    def saveFile(self):
        options = QFileDialog.Options()
        file_format, _ = QFileDialog.getSaveFileName(
            self, "Сохранить файл", "", "Text Files (*.txt);;Word Documents (*.doc)", options=options)
        if file_format:
            text = self.text_edit.toPlainText()
            if file_format.endswith('.txt'):
                self.saveTextFile(file_format, text)
            elif file_format.endswith('.doc'):
                self.saveDocFile(file_format, text)

    def saveTextFile(self, filename, text):
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(text)

    def saveDocFile(self, filename, text):
        doc = docx.Document()
        doc.add_paragraph(text)
        doc.save(filename)

    def show_actions_for_day(self):
        actions_for_day = self.get_actions_for_day()
        self.text_edit.setPlainText('\n'.join(actions_for_day))

    def show_actions_for_week(self):
        actions_for_week = self.get_actions_for_week()
        self.text_edit.setPlainText('\n'.join(actions_for_week))

    def show_actions_for_month(self):
        actions_for_month = self.get_actions_for_month()
        self.text_edit.setPlainText('\n'.join(actions_for_month))

    def get_actions_for_day(self):
        today = datetime.now().strftime('%Y-%m-%d')
        self.c.execute("SELECT category, roberta_category, file_path FROM actions WHERE date LIKE ?", (f'{today}%',))
        actions_for_day = self.c.fetchall()
        return [f"Category: {action[0]}, RoBERTa Category: {action[1]}, File: {action[2]}" for action in actions_for_day]

    def get_actions_for_week(self):
        today = datetime.now()
        start_of_week = today - timedelta(days=today.weekday())
        start_of_week_str = start_of_week.strftime('%Y-%m-%d')
        end_of_week_str = (start_of_week + timedelta(days=6)).strftime('%Y-%m-%d')
        self.c.execute("SELECT category, roberta_category, file_path FROM actions WHERE date BETWEEN ? AND ?",
                       (start_of_week_str, end_of_week_str))
        actions_for_week = self.c.fetchall()
        return [f"Category: {action[0]}, RoBERTa Category: {action[1]}, File: {action[2]}" for action in actions_for_week]

    def get_actions_for_month(self):
        today = datetime.now()
        start_date = today - timedelta(days=30)
        start_date_str = start_date.strftime('%Y-%m-%d')
        end_date_str = today.strftime('%Y-%m-%d')
        self.c.execute("SELECT category, roberta_category, file_path FROM actions WHERE date BETWEEN ? AND ?",
                       (start_date_str, end_date_str))
        actions_for_month = self.c.fetchall()
        return [f"Category: {action[0]}, RoBERTa Category: {action[1]}, File: {action[2]}" for action in actions_for_month]

    def displayContent(self, filename):
        text = ''
        if filename.endswith('.docx'):
            text = self.extractTextFromDocx(filename)
        elif filename.endswith('.pdf'):
            text = self.extractTextFromPDF(filename)
        elif filename.endswith(('.png', '.jpg', '.jpeg')):
            text = self.extractTextFromImage(filename)
        self.text_edit.setPlainText(text)
        if text.strip():
            category_name, roberta_category_name = self.classifyText(text)
            self.label_category.setText(f'Категория: {category_name}')
            self.label_roberta_category.setText(f'Категория (RoBERTa): {roberta_category_name}')
            self.save_action(filename, category_name, roberta_category_name)
        else:
            self.label_category.setText('Категория: Текст не найден')
            self.label_roberta_category.setText('Категория (RoBERTa): Текст не найден')

    def save_action(self, file_path, category, roberta_category):
        current_date = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        self.c.execute("INSERT INTO actions (date, category, roberta_category, file_path) VALUES (?, ?, ?, ?)",
                       (current_date, category, roberta_category, file_path))
        self.conn.commit()

    def extractTextFromDocx(self, filename):
        doc = docx.Document(filename)
        full_text = [para.text for para in doc.paragraphs]
        return '\n'.join(full_text)

    def extractTextFromPDF(self, filename):
        with pdfplumber.open(filename) as pdf:
            full_text = [page.extract_text() for page in pdf.pages if page.extract_text() is not None]
        return '\n'.join(full_text)

    def extractTextFromImage(self, filename):
        image = Image.open(filename)
        text = pytesseract.image_to_string(image, lang='rus')
        return text

    def classifyText(self, text):
        text_vector = self.vectorizer.transform([text])
        prediction = self.model.predict(text_vector)
        category_name = self.number_to_category[prediction[0]]

        # Получаем максимальную длину последовательности из конфигурации модели
        max_length = self.roberta_model.config.max_position_embeddings
        
        # Токенизируем текст с учетом максимальной длины и добавляем padding
        inputs = self.tokenizer(text, truncation=True, max_length=max_length, padding="max_length", return_tensors="pt")
        
        # Получаем маски внимания из входных данных
        attention_mask = inputs["attention_mask"]
        
        # Получаем position_ids
        position_ids = torch.arange(inputs["input_ids"].size(1)).expand(1, -1)
        
        # Переносим входные данные на устройство модели
        inputs = {k: v.to(self.roberta_model.device) for k, v in inputs.items()}
        position_ids = position_ids.to(self.roberta_model.device)
        
        # Передаем маски внимания в модель
        outputs = self.roberta_model(input_ids=inputs["input_ids"], attention_mask=attention_mask, position_ids=position_ids)
        
        logits = outputs.logits
        roberta_prediction = torch.argmax(logits, dim=1).item()
        try:
            roberta_category_name = self.roberta_number_to_category[roberta_prediction]
        except KeyError:
            roberta_category_name = "Unknown"

        return category_name, roberta_category_name

if __name__ == '__main__':
    app = QApplication(sys.argv)
    ex = TextExtractorApp()
    ex.show()
    sys.exit(app.exec_())